from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, List
import logging

logger = logging.getLogger(__name__)

class DocxGenerator:
    def __init__(self):
        self.document = Document()
        self.setup_styles()

    def setup_styles(self):
        # Set default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

    def add_heading(self, text, level=1):
        heading = self.document.add_heading(text, level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return heading

    def add_paragraph(self, text, style=None):
        paragraph = self.document.add_paragraph(text, style)
        return paragraph

    def add_bullet_points(self, items):
        for item in items:
            paragraph = self.document.add_paragraph(style='List Bullet')
            paragraph.add_run(item)

    def generate_resume(self, resume_data, output_path):
        try:
            # Clear any existing content
            self.document = Document()
            self.setup_styles()

            # Add name and title with enhanced formatting
            name = resume_data.get('name', '')
            title = resume_data.get('title', '')
            if name:
                name_paragraph = self.document.add_paragraph()
                name_run = name_paragraph.add_run(name)
                name_run.bold = True
                name_run.font.size = Pt(24)  # Increased font size to 24pt
                name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if title:
                title_paragraph = self.document.add_paragraph()
                title_run = title_paragraph.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(14)  # Increased font size to 14pt
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add contact information
            contact_info = []
            if resume_data.get('phone'):
                contact_info.append(resume_data['phone'])
            if resume_data.get('email'):
                contact_info.append(resume_data['email'])
            if resume_data.get('location'):
                contact_info.append(resume_data['location'])
            if contact_info:
                contact_paragraph = self.document.add_paragraph()
                contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_paragraph.add_run(" | ".join(contact_info))

            # Add summary
            summary = resume_data.get('summary', '')
            if summary:
                self.add_heading('Summary', level=1)
                self.add_paragraph(summary)

            # Add education
            education = resume_data.get('education', [])
            if education:
                self.add_heading('Education', level=1)
                for edu in education:
                    self._add_education(self.document, edu)

            # Add skills
            skills = resume_data.get('skills', [])
            if skills:
                self.add_heading('Skills', level=1)
                self.add_paragraph(", ".join(skills))

            # Add projects
            projects = resume_data.get('projects', [])
            if projects:
                self.add_heading('Projects', level=1)
                for project in projects:
                    if project.get('name'):
                        self.add_paragraph(project['name'], style='Heading 2')
                    if project.get('description'):
                        self.add_paragraph(project['description'])

            # Add experience
            experience = resume_data.get('experience', [])
            if experience:
                self.add_heading('Experience', level=1)
                for exp in experience:
                    exp_text = []
                    if exp.get('title'):
                        exp_text.append(exp['title'])
                    if exp.get('company'):
                        exp_text.append(f"at {exp['company']}")
                    if exp.get('duration'):
                        exp_text.append(f"({exp['duration']})")
                    if exp_text:
                        self.add_paragraph(" ".join(exp_text), style='Heading 2')
                    if exp.get('description'):
                        self.add_paragraph(exp['description'])

            # Add certifications
            certifications = resume_data.get('certifications', [])
            if certifications:
                self.add_heading('Certifications', level=1)
                for cert in certifications:
                    cert_text = []
                    if cert.get('title'):
                        cert_text.append(cert['title'])
                    if cert.get('issuer'):
                        cert_text.append(f"- {cert['issuer']}")
                    if cert.get('date'):
                        cert_text.append(f"({cert['date']})")
                    if cert_text:
                        self.add_paragraph(" ".join(cert_text))

            # Save the document
            self.document.save(output_path)
            logger.info(f"Resume saved to {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            raise

    def _add_education(self, doc, education_data):
        """Add education section to the document."""
        p = doc.add_paragraph()
        p.add_run(education_data.get('institution', '')).bold = True
        p.add_run(f" - {education_data.get('degree', '')}")
        
        # Add year range and CGPA if available
        details = []
        if education_data.get('year_range'):
            details.append(f"Year: {education_data.get('year_range')}")
        if education_data.get('cgpa'):
            details.append(f"CGPA: {education_data.get('cgpa')}")
        if education_data.get('location'):
            details.append(f"Location: {education_data.get('location')}")
        
        if details:
            p = doc.add_paragraph()
            p.add_run(" | ".join(details)).italic = True 